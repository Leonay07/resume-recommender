import os
import re
from typing import Dict, List, Optional, Tuple
import logging

# PDF parsing
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: pdfplumber not installed. PDF parsing will not be available.")

# DOCX parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX parsing will not be available.")

# Import skill dictionary
try:
    from .skills_dict import get_all_skills, normalize_skill, SKILL_DICT
except ImportError:
    # For standalone testing
    from skills_dict import get_all_skills, normalize_skill, SKILL_DICT

# ========================================
# Logging Configuration
# ========================================

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ========================================
# Constants and Keywords
# ========================================

# Section keywords for resume parsing
# Extended with more variations and common formats
SECTION_KEYWORDS = {
    'skills': [
        'skills', 'skill', 'technical skills', 'core competencies',
        'technologies', 'expertise', 'proficiencies',
        'technical expertise', 'programming skills',
        'tools and technologies', 'technical proficiencies',
        'competencies', 'technical competencies', 'areas of expertise',
        'technical summary', 'technology skills', 'key skills',
        'professional skills', 'software skills', 'computer skills',
        'qualifications', 'technical qualifications', 'abilities',
        'certifications and skills', 'skills and certifications',
        'languages and tools', 'tools', 'programming languages',
        'frameworks', 'software', 'platforms'
    ],
    'experience': [
        'experience', 'work experience', 'professional experience',
        'employment history', 'work history', 'career history',
        'professional background', 'employment', 'career',
        'job experience', 'relevant experience', 'professional history',
        'work summary', 'employment experience', 'positions held',
        'career experience', 'job history', 'internship',
        'internships', 'internship experience', 'co-op experience'
    ],
    'education': [
        'education', 'academic background', 'qualifications',
        'academic credentials', 'educational background',
        'degrees', 'academic history', 'academic',
        'educational qualifications', 'schooling', 'training',
        'academic qualifications', 'certifications', 'credentials',
        'education and certifications', 'courses', 'coursework'
    ],
    'projects': [
        'projects', 'personal projects', 'academic projects',
        'project experience', 'selected projects', 'key projects',
        'notable projects', 'portfolio', 'side projects',
        'professional projects', 'relevant projects', 'project work',
        'technical projects', 'research projects', 'capstone',
        'capstone project', 'thesis', 'dissertation'
    ],
    'summary': [
        'summary', 'professional summary', 'objective',
        'profile', 'about me', 'career objective',
        'professional profile', 'executive summary',
        'career summary', 'personal statement',
        'about', 'introduction', 'bio', 'overview',
        'professional objective', 'goals', 'career goals',
        'statement', 'highlights', 'professional highlights'
    ]
}

# Job role keywords for inference
JOB_ROLE_KEYWORDS = {
    'Data Scientist': [
        'data scientist', 'data science', 'machine learning',
        'statistical modeling', 'data analysis', 'predictive modeling',
        'data mining', 'statistics', 'quantitative analysis',
        'statistical analysis', 'data analytics'
    ],
    'Machine Learning Engineer': [
        'machine learning engineer', 'ml engineer', 'deep learning',
        'neural networks', 'model deployment', 'mlops',
        'ai engineer', 'model training', 'ml systems',
        'deep learning engineer'
    ],
    'Data Engineer': [
        'data engineer', 'data engineering', 'data pipeline',
        'etl', 'data warehouse', 'big data', 'spark', 'hadoop',
        'data infrastructure', 'data architecture',
        'pipeline development'
    ],
    'Data Analyst': [
        'data analyst', 'business analyst', 'analytics',
        'reporting', 'dashboard', 'sql analyst',
        'data visualization', 'business intelligence',
        'analyst', 'bi analyst'
    ],
    'Software Engineer': [
        'software engineer', 'software developer', 'full stack',
        'backend developer', 'frontend developer', 'web developer',
        'application developer', 'software development',
        'full-stack developer', 'backend engineer'
    ],
    'DevOps Engineer': [
        'devops', 'devops engineer', 'site reliability',
        'cloud engineer', 'infrastructure', 'ci/cd',
        'kubernetes', 'sre', 'platform engineer',
        'infrastructure engineer'
    ],
    'Research Scientist': [
        'research scientist', 'researcher', 'phd', 'publications',
        'academic research', 'scientific research',
        'research engineer', 'scientist'
    ],
    'AI Engineer': [
        'ai engineer', 'artificial intelligence', 'ai/ml',
        'ai systems', 'ai developer', 'ai researcher',
        'artificial intelligence engineer'
    ]
}

# ========================================
# Helper Functions
# ========================================

def _clean_text(text: str) -> str:
    """
    Clean and normalize text.

    Args:
        text: Raw text to clean

    Returns:
        Cleaned text
    """
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters but keep basic punctuation
    text = re.sub(r'[^\w\s.,;:!?()\-/\'\"]', '', text)
    return text.strip()


def _is_section_header(line: str, max_length: int = 80) -> bool:
    """
    Check if a line is likely a section header.

    Enhanced detection with multiple heuristics to handle various resume formats:
    - ALL CAPS headers (e.g., "SKILLS", "EXPERIENCE")
    - Title Case headers (e.g., "Skills", "Work Experience")
    - Headers with colons (e.g., "Skills:", "Experience:")
    - Headers with special markers (e.g., "--- Skills ---", "== Experience ==")
    - Short lines that could be headers

    Args:
        line: Text line to check
        max_length: Maximum length for a header (increased to 80)

    Returns:
        True if line appears to be a header
    """
    line = line.strip()

    # Skip empty or very short lines
    if len(line) < 2:
        return False

    # Skip lines that are too long to be headers
    if len(line) > max_length:
        return False

    # Remove common decorative characters for analysis
    clean_line = re.sub(r'^[\s\-=_•*#|:]+|[\s\-=_•*#|:]+$', '', line).strip()

    # Skip if nothing left after cleaning
    if len(clean_line) < 2:
        return False

    # Check 1: ALL CAPS (strong indicator)
    if clean_line.isupper() and len(clean_line) >= 2:
        return True

    # Check 2: Ends with colon (common header format)
    if line.endswith(':'):
        return True

    # Check 3: Starts with common decorative markers
    if re.match(r'^[\-=_•*#]{2,}', line) or re.match(r'^[|►▶→●○◆◇■□▪▫]+', line):
        return True

    # Check 4: Title Case with limited words (likely header)
    words = clean_line.split()
    word_count = len(words)

    if word_count <= 5:
        # Check if it looks like a title (starts with uppercase, few words)
        if clean_line[0].isupper():
            # Additional check: headers usually don't contain certain patterns
            # Skip lines that look like sentences or bullet points
            if not re.search(r'\d{4}', clean_line):  # No years (like "2020-2023")
                if not re.match(r'^[\-•*]\s', line):  # Not a bullet point
                    if word_count <= 4:
                        return True

    # Check 5: Line is wrapped in decorative characters
    if re.match(r'^[\-=_*#]+\s*\w+.*\s*[\-=_*#]+$', line):
        return True

    # Check 6: Very short lines (1-3 words) that start with uppercase
    if word_count <= 3 and clean_line[0].isupper():
        # Skip if it looks like a name (two capitalized words)
        if not (word_count == 2 and all(w[0].isupper() for w in words)):
            return True

    return False


def _match_section_keyword(line: str, section_keywords: dict) -> Optional[str]:
    """
    Match a line against section keywords.

    Enhanced matching with:
    - Case-insensitive matching
    - Cleaned line matching (removes decorations)
    - Partial word matching for common headers

    Args:
        line: The line to check
        section_keywords: Dictionary of section names to keyword lists

    Returns:
        Section name if matched, None otherwise
    """
    line_lower = line.lower().strip()

    # Remove common decorative characters
    clean_line = re.sub(r'^[\s\-=_•*#|:]+|[\s\-=_•*#|:]+$', '', line_lower).strip()
    # Also remove trailing colons
    clean_line = clean_line.rstrip(':').strip()

    for section_name, keywords in section_keywords.items():
        for keyword in keywords:
            keyword_lower = keyword.lower()
            # Exact match after cleaning
            if clean_line == keyword_lower:
                return section_name
            # Keyword is contained in the line
            if keyword_lower in clean_line:
                return section_name
            # Line is contained in the keyword (for abbreviated headers)
            if len(clean_line) >= 3 and clean_line in keyword_lower:
                return section_name

    return None


# ========================================
# Main Resume Parser Class
# ========================================

class ResumeParser:
    """
    Main resume parser class for extracting information from resumes.

    This class handles:
    - Loading resume files (PDF/DOCX)
    - Parsing resume sections
    - Extracting skills
    - Inferring target job roles
    """

    def __init__(self):
        """Initialize the resume parser."""
        self.all_skills = get_all_skills()
        logger.info("ResumeParser initialized with %d skills", len(self.all_skills))

    # ========================================
    # File Loading Methods
    # ========================================

    def load_resume(self, file_path: str) -> str:
        """
        Load a resume file and extract its text content.

        Supports PDF and DOCX formats. Automatically detects file type
        based on extension.

        Args:
            file_path: Path to the resume file

        Returns:
            Extracted text content

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is not supported

        Example:
            >>> parser = ResumeParser()
            >>> text = parser.load_resume("resume.pdf")
            >>> print(len(text))
            5000
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        _, file_extension = os.path.splitext(file_path)
        file_extension = file_extension.lower()

        if file_extension == '.pdf':
            return self._load_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self._load_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    def _load_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text
        """
        if not PDF_AVAILABLE:
            raise ImportError("pdfplumber is required for PDF parsing. Install with: pip install pdfplumber")

        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            logger.info(f"PDF parsed successfully: {len(text)} characters from {file_path}")
            return text.strip()

        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            raise

    def _load_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

        try:
            doc = Document(file_path)
            text = ""

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Extract text from tables (if any)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            logger.info(f"DOCX parsed successfully: {len(text)} characters from {file_path}")
            return text.strip()

        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise

    # ========================================
    # Section Parsing Methods
    # ========================================

    def parse_sections(self, resume_text: str) -> Dict[str, str]:
        """
        Parse resume text into different sections.

        Enhanced parsing with multiple strategies:
        1. First pass: Look for explicit section headers
        2. Second pass: If no sections found, try keyword-based detection
        3. Fallback: Keep all content accessible for skill extraction

        Identifies common resume sections like Skills, Experience, Education,
        Projects, and Summary based on keyword matching.

        Args:
            resume_text: Complete resume text

        Returns:
            Dictionary mapping section names to their content

        Example:
            >>> parser = ResumeParser()
            >>> sections = parser.parse_sections(resume_text)
            >>> print(sections.keys())
            dict_keys(['skills', 'experience', 'education', 'projects', 'summary'])
        """
        sections = {
            'skills': '',
            'experience': '',
            'education': '',
            'projects': '',
            'summary': '',
            'other': ''
        }

        lines = resume_text.split('\n')
        current_section = 'other'
        sections_found = []

        # First pass: standard header-based parsing
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue

            # Check if this line is a section header
            section_found = False

            if _is_section_header(line_stripped):
                # Use enhanced keyword matching
                matched_section = _match_section_keyword(line_stripped, SECTION_KEYWORDS)
                if matched_section:
                    current_section = matched_section
                    section_found = True
                    if matched_section not in sections_found:
                        sections_found.append(matched_section)
                    logger.debug(f"Found section header: '{line_stripped}' -> {matched_section}")

            # Add content to current section (skip the header line itself)
            if not section_found:
                sections[current_section] += line + '\n'

        # Clean up sections
        for key in sections:
            sections[key] = sections[key].strip()

        # Second pass: If no meaningful sections found, try alternative detection
        meaningful_sections = [s for s in sections_found if s != 'other']

        if not meaningful_sections:
            logger.warning("No section headers detected. Attempting keyword-based section detection...")
            sections = self._fallback_section_detection(resume_text)
            sections_found = [k for k, v in sections.items() if v and k != 'other']

        # Log section statistics
        section_counts = {k: len(v) for k, v in sections.items() if v}
        logger.info(f"Sections parsed: {list(section_counts.keys())}")

        if not meaningful_sections:
            logger.warning("Could not identify distinct sections. "
                          "Skills will be extracted from full resume text.")

        return sections

    def _fallback_section_detection(self, resume_text: str) -> Dict[str, str]:
        """
        Fallback method to detect sections when header-based parsing fails.

        Uses keyword proximity and content analysis to identify sections
        even without clear headers.

        Args:
            resume_text: Complete resume text

        Returns:
            Dictionary mapping section names to their content
        """
        sections = {
            'skills': '',
            'experience': '',
            'education': '',
            'projects': '',
            'summary': '',
            'other': ''
        }

        lines = resume_text.split('\n')
        full_text = resume_text.lower()

        # Strategy 1: Look for inline section indicators
        # Some resumes have "Skills: Python, Java, ..." format
        inline_patterns = {
            'skills': [
                r'(?:skills|technologies|tools|languages)\s*[:\-]\s*(.+)',
                r'(?:proficient\s+(?:in|with)|experienced\s+(?:in|with))\s*[:\-]?\s*(.+)',
            ],
            'education': [
                r'(?:university|college|bachelor|master|phd|degree)\s*[:\-]?\s*(.+)',
            ]
        }

        for section_name, patterns in inline_patterns.items():
            for pattern in patterns:
                matches = re.findall(pattern, full_text, re.IGNORECASE | re.MULTILINE)
                if matches:
                    sections[section_name] += ' '.join(matches) + '\n'

        # Strategy 2: Segment by keywords in the text
        # Look for lines that contain section keywords
        current_section = 'other'
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue

            line_lower = line_stripped.lower()

            # Check if line contains a section keyword
            for section_name, keywords in SECTION_KEYWORDS.items():
                for keyword in keywords:
                    # Check if keyword appears at start of line or standalone
                    if (line_lower.startswith(keyword) or
                        re.match(rf'^[\W]*{re.escape(keyword)}[\W]*$', line_lower) or
                        (keyword in line_lower and len(line_stripped) < 50)):
                        current_section = section_name
                        break
                else:
                    continue
                break

            sections[current_section] += line + '\n'

        # Clean up
        for key in sections:
            sections[key] = sections[key].strip()

        # If still no skills section found, put all content in 'other'
        # This ensures extract_skills can search the full text
        if not sections['skills'] and not any(sections[k] for k in ['experience', 'education', 'projects']):
            sections['other'] = resume_text

        return sections

    # ========================================
    # Skill Extraction Methods
    # ========================================

    def extract_skills(self, sections: Dict[str, str]) -> Dict[str, any]:
        """
        Extract technical skills from resume sections.

        Uses the skill dictionary to identify skills mentioned in the resume.
        Distinguishes between primary skills (found in Skills section) and
        secondary skills (found in other sections).

        Enhanced with fallback logic:
        - If no skills found in designated sections, searches 'other' section
        - Ensures skills are always extracted even if section parsing fails

        Args:
            sections: Dictionary of resume sections from parse_sections()

        Returns:
            Dictionary containing:
            - primary_skills: Skills found in Skills section
            - secondary_skills: Skills found in other sections
            - all_skills: All unique skills found
            - skill_frequency: Count of how many times each skill appears
            - total_count: Total number of unique skills

        Example:
            >>> parser = ResumeParser()
            >>> sections = parser.parse_sections(resume_text)
            >>> skills = parser.extract_skills(sections)
            >>> print(skills['all_skills'])
            ['Python', 'TensorFlow', 'AWS', 'Docker']
        """
        primary_skills = set()      # Skills in Skills section
        secondary_skills = set()    # Skills in other sections
        skill_frequency = {}        # Track how often each skill appears

        # Search in Skills section first
        skills_section = sections.get('skills', '')
        for skill in self.all_skills:
            # Use word boundary matching to avoid partial matches
            pattern = r'\b' + re.escape(skill) + r'\b'
            matches = re.findall(pattern, skills_section, re.IGNORECASE)

            if matches:
                normalized = normalize_skill(skill)
                primary_skills.add(normalized)
                skill_frequency[normalized] = len(matches)

        # Search in other named sections (experience, projects, summary, education)
        named_sections = (
            sections.get('experience', '') + ' ' +
            sections.get('projects', '') + ' ' +
            sections.get('summary', '') + ' ' +
            sections.get('education', '')
        )

        for skill in self.all_skills:
            pattern = r'\b' + re.escape(skill) + r'\b'
            matches = re.findall(pattern, named_sections, re.IGNORECASE)

            if matches:
                normalized = normalize_skill(skill)
                if normalized not in primary_skills:
                    secondary_skills.add(normalized)
                    skill_frequency[normalized] = len(matches)
                else:
                    # Add to existing count
                    skill_frequency[normalized] += len(matches)

        # CRITICAL FALLBACK: If no skills found yet, search in 'other' section
        # This handles cases where section parsing failed and all content is in 'other'
        all_found = primary_skills.union(secondary_skills)

        if not all_found:
            other_section = sections.get('other', '')
            if other_section:
                logger.info("No skills found in named sections. "
                           "Searching in 'other' section as fallback...")

                for skill in self.all_skills:
                    pattern = r'\b' + re.escape(skill) + r'\b'
                    matches = re.findall(pattern, other_section, re.IGNORECASE)

                    if matches:
                        normalized = normalize_skill(skill)
                        # Treat all skills from 'other' as secondary since
                        # we couldn't identify the Skills section
                        secondary_skills.add(normalized)
                        skill_frequency[normalized] = skill_frequency.get(normalized, 0) + len(matches)

                if secondary_skills:
                    logger.info(f"Fallback extraction found {len(secondary_skills)} skills "
                               "from 'other' section")

        # Even if we found some skills, also search 'other' to ensure completeness
        # This catches cases where some content wasn't properly categorized
        elif sections.get('other', ''):
            other_section = sections.get('other', '')
            additional_found = 0

            for skill in self.all_skills:
                pattern = r'\b' + re.escape(skill) + r'\b'
                matches = re.findall(pattern, other_section, re.IGNORECASE)

                if matches:
                    normalized = normalize_skill(skill)
                    if normalized not in primary_skills and normalized not in secondary_skills:
                        secondary_skills.add(normalized)
                        skill_frequency[normalized] = len(matches)
                        additional_found += 1
                    elif normalized in skill_frequency:
                        skill_frequency[normalized] += len(matches)

            if additional_found > 0:
                logger.debug(f"Found {additional_found} additional skills in 'other' section")

        all_skills = primary_skills.union(secondary_skills)

        result = {
            'primary_skills': sorted(list(primary_skills)),
            'secondary_skills': sorted(list(secondary_skills)),
            'all_skills': sorted(list(all_skills)),
            'skill_frequency': skill_frequency,
            'total_count': len(all_skills)
        }

        logger.info(f"Extracted {result['total_count']} skills "
                   f"({len(primary_skills)} primary, {len(secondary_skills)} secondary)")

        return result

    # ========================================
    # Job Role Inference Methods
    # ========================================

    def infer_roles(self, sections: Dict[str, str],
                    user_input: Optional[str] = None) -> List[str]:
        """
        Infer target job roles from resume content.

        Analyzes the resume to determine what job roles the candidate is
        targeting. Considers explicit user input if provided, and falls back
        to analyzing summary and experience sections.

        Enhanced with fallback logic:
        - Searches 'other' section if named sections are empty
        - Ensures role inference works even if section parsing fails

        Args:
            sections: Dictionary of resume sections
            user_input: Optional explicit job title from user

        Returns:
            List of inferred job roles (up to 3)

        Example:
            >>> parser = ResumeParser()
            >>> sections = parser.parse_sections(resume_text)
            >>> roles = parser.infer_roles(sections, user_input="Data Scientist")
            >>> print(roles)
            ['Data Scientist', 'Machine Learning Engineer']
        """
        target_roles = []

        # Step 1: Handle explicit user input
        if user_input:
            user_input_clean = user_input.strip()

            # Check if user input matches any of our predefined roles
            for role in JOB_ROLE_KEYWORDS.keys():
                if (user_input_clean.lower() in role.lower() or
                    role.lower() in user_input_clean.lower()):
                    if role not in target_roles:
                        target_roles.append(role)

            # If no match found, add user input directly
            if not target_roles:
                target_roles.append(user_input_clean)
                logger.info(f"Using custom role from user: {user_input_clean}")

        # Step 2: Analyze Summary/Objective section
        summary_text = sections.get('summary', '').lower()
        role_scores = {}

        if summary_text:
            for role, keywords in JOB_ROLE_KEYWORDS.items():
                score = 0
                for keyword in keywords:
                    if keyword in summary_text:
                        score += 1
                if score > 0:
                    role_scores[role] = score

        # Step 3: If no summary info, analyze skills and experience
        if not role_scores:
            skills_text = sections.get('skills', '').lower()
            experience_text = sections.get('experience', '').lower()
            combined_text = skills_text + ' ' + experience_text

            for role, keywords in JOB_ROLE_KEYWORDS.items():
                score = 0
                for keyword in keywords:
                    if keyword in combined_text:
                        score += 1
                if score > 0:
                    role_scores[role] = score

        # Step 4: FALLBACK - If still no roles found, search in 'other' section
        # This handles cases where section parsing failed
        if not role_scores:
            other_text = sections.get('other', '').lower()
            if other_text:
                logger.info("No roles found in named sections. "
                           "Searching in 'other' section as fallback...")

                for role, keywords in JOB_ROLE_KEYWORDS.items():
                    score = 0
                    for keyword in keywords:
                        if keyword in other_text:
                            score += 1
                    if score > 0:
                        role_scores[role] = score

                if role_scores:
                    logger.info(f"Fallback found {len(role_scores)} potential roles "
                               "from 'other' section")

        # Sort roles by score and add top 3 (excluding user input if already added)
        sorted_roles = sorted(role_scores.items(),
                            key=lambda x: x[1], reverse=True)

        for role, score in sorted_roles[:3]:
            if role not in target_roles:
                target_roles.append(role)

        # If still no roles found, return a default
        if not target_roles:
            target_roles = ['General']
            logger.warning("No specific role could be inferred, using 'General'")

        logger.info(f"Inferred roles: {target_roles}")
        return target_roles

    # ========================================
    # Main Parsing Method
    # ========================================

    def parse(self, file_path: str,
              user_title: Optional[str] = None) -> Dict:
        """
        Main method to parse a resume file and extract all information.

        This is the primary interface for the resume parser. It orchestrates
        all parsing steps: loading the file, identifying sections, extracting
        skills, and inferring roles.

        Args:
            file_path: Path to resume file (PDF or DOCX)
            user_title: Optional job title specified by user

        Returns:
            Dictionary containing:
            - success: Boolean indicating if parsing succeeded
            - resume_text: Full text of resume
            - sections: Dictionary of identified sections
            - skills: Dictionary of extracted skills
            - target_roles: List of inferred job roles
            - metadata: File information
            - error: Error message if parsing failed

        Example:
            >>> parser = ResumeParser()
            >>> result = parser.parse("resume.pdf", user_title="Data Scientist")
            >>> if result['success']:
            ...     print(f"Found {result['skills']['total_count']} skills")
            ...     print(f"Target roles: {result['target_roles']}")
        """
        try:
            logger.info(f"Starting resume parsing: {file_path}")

            # Step 1: Load resume text
            resume_text = self.load_resume(file_path)

            if not resume_text:
                raise ValueError("No text could be extracted from the file")

            # Step 2: Parse sections
            sections = self.parse_sections(resume_text)

            # Step 3: Extract skills
            skills_result = self.extract_skills(sections)

            # Step 4: Infer target roles
            target_roles = self.infer_roles(sections, user_title)

            # Step 5: Compile results
            result = {
                'success': True,
                'resume_text': resume_text,
                'sections': sections,
                'skills': skills_result,
                'target_roles': target_roles,
                'metadata': {
                    'file_path': file_path,
                    'file_type': os.path.splitext(file_path)[1],
                    'text_length': len(resume_text),
                    'word_count': len(resume_text.split())
                },
                'error': None
            }

            logger.info("Resume parsing completed successfully")
            return result

        except Exception as e:
            error_msg = f"Error parsing resume: {str(e)}"
            logger.error(error_msg)

            return {
                'success': False,
                'error': error_msg,
                'resume_text': None,
                'sections': None,
                'skills': None,
                'target_roles': None,
                'metadata': {
                    'file_path': file_path,
                    'error_type': type(e).__name__
                }
            }


# ========================================
# Convenience Functions
# ========================================

def parse_resume(file_path: str,
                 user_title: Optional[str] = None) -> Dict:
    """
    Convenience function to parse a resume.

    This is a simplified interface that creates a ResumeParser instance
    and calls its parse method.

    Args:
        file_path: Path to resume file
        user_title: Optional user-specified job title

    Returns:
        Dictionary with parsing results

    Example:
        >>> result = parse_resume("resume.pdf", user_title="Data Scientist")
        >>> print(result['skills']['all_skills'])
    """
    parser = ResumeParser()
    return parser.parse(file_path, user_title)


# ========================================
# Module Testing
# ========================================

if __name__ == "__main__":
    print("=" * 70)
    print("RESUME PARSER MODULE - STANDALONE TEST")
    print("=" * 70)

    # Test with sample text
    sample_resume = """
    John Doe
    Data Scientist

    PROFESSIONAL SUMMARY
    Experienced data scientist with expertise in machine learning and Python.

    SKILLS
    Python, TensorFlow, PyTorch, scikit-learn, AWS, Docker, Kubernetes
    SQL, MongoDB, Spark, Pandas, NumPy

    EXPERIENCE
    Senior Data Scientist at Google (2020-2023)
    - Developed machine learning models using TensorFlow and PyTorch
    - Built data pipelines with Apache Spark
    - Deployed models on AWS using Docker and Kubernetes

    EDUCATION
    Master of Science in Computer Science
    Stanford University (2018-2020)

    PROJECTS
    - Built a recommendation system using collaborative filtering
    - Created a NLP model for sentiment analysis
    """

    print("\n[TEST] Creating ResumeParser instance...")
    parser = ResumeParser()
    print(f"✓ Parser initialized with {len(parser.all_skills)} skills")

    print("\n[TEST] Parsing sample resume text...")
    sections = parser.parse_sections(sample_resume)
    print(f"✓ Found sections: {[k for k, v in sections.items() if v]}")

    print("\n[TEST] Extracting skills...")
    skills = parser.extract_skills(sections)
    print(f"✓ Extracted {skills['total_count']} skills:")
    print(f"  Primary: {skills['primary_skills'][:5]}...")
    print(f"  Secondary: {skills['secondary_skills'][:5]}...")

    print("\n[TEST] Inferring target roles...")
    roles = parser.infer_roles(sections)
    print(f"✓ Inferred roles: {roles}")

    print("\n" + "=" * 70)
    print("STANDALONE TESTS COMPLETED SUCCESSFULLY ✓")
    print("=" * 70)
    print("\nNote: To test with actual PDF/DOCX files, use:")
    print("  result = parse_resume('path/to/resume.pdf')")
    print("  print(result['skills']['all_skills'])")


# ========================================
# Division.md Interface Functions
# ========================================

def extract_resume_skills(resume_sections, skill_dict=None):
    """
    Extract skills from resume sections.

    This is the interface function specified in Division.md.
    It wraps the ResumeParser.extract_skills() method.

    Args:
        resume_sections (dict): Resume sections from parse_sections()
        skill_dict (dict, optional): Custom skill dictionary.
                                     If None, uses default SKILL_DICT

    Returns:
        dict: {
            'all_skills': list,
            'primary_skills': list,
            'secondary_skills': list,
            'skill_frequency': dict,
            'total_count': int
        }

    Example:
        >>> from nlp_model import load_resume, extract_resume_skills
        >>> sections = parser.parse_sections(resume_text)
        >>> skills = extract_resume_skills(sections)
        >>> print(skills['all_skills'])
    """
    parser = ResumeParser()

    if skill_dict:
        try:
            if isinstance(skill_dict, dict):
                custom_skills = []
                for value in skill_dict.values():
                    if isinstance(value, (list, tuple, set)):
                        custom_skills.extend(value)
                    elif isinstance(value, str):
                        custom_skills.append(value)
            elif isinstance(skill_dict, (list, tuple, set)):
                custom_skills = list(skill_dict)
            else:
                raise TypeError("skill_dict must be a dict or iterable of strings")

            normalized = {normalize_skill(skill) for skill in custom_skills if isinstance(skill, str)}
            if normalized:
                parser.all_skills = sorted(normalized)
                logger.debug(
                    "extract_resume_skills using custom skill dictionary (%d skills)",
                    len(normalized)
                )

        except Exception as exc:
            logger.warning(
                "Failed to apply custom skill_dict (%s). Falling back to default SKILL_DICT.",
                exc
            )

    return parser.extract_skills(resume_sections)


def infer_target_roles(resume_sections, optional_user_input=None):
    """
    Infer target job roles from resume sections.

    This is the interface function specified in Division.md.
    It wraps the ResumeParser.infer_roles() method.

    Based on summary/objective sections and supports explicit user input.

    Args:
        resume_sections (dict): Resume sections from parse_sections()
        optional_user_input (str, optional): User-specified target role

    Returns:
        list: List of inferred target roles

    Example:
        >>> from nlp_model import infer_target_roles
        >>> sections = parser.parse_sections(resume_text)
        >>> roles = infer_target_roles(sections, "Data Scientist")
        >>> print(roles)  # ['Data Scientist', 'Data Analyst', ...]
    """
    parser = ResumeParser()
    return parser.infer_roles(resume_sections, user_input=optional_user_input)
